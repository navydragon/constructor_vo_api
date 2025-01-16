from django.http import HttpResponse
from docx import Document
from docx.shared import Pt

from io import BytesIO
from urllib.parse import quote

from competenceprofile.models import Knowledge, Ability
from disciplines.models import Discipline
from products.models import Product
from programs.models import Program
from django.shortcuts import get_object_or_404
from django.db.models import Count


def export_design(request, program_id):
    # Создаем docx файл
    doc = Document()

    # программа
    program = get_object_or_404(Program, id=program_id)
    doc.add_heading(f'Дизайн программы «{program.direction_id.code} {program.direction_id.name} {program.profile}»', 0)
    doc.add_heading('Реконструкция деятельности',1)
    doc.add_heading('Продукты', 2)

    for index,product in enumerate( program.products.all()):
        pr_number = index + 1
        doc.add_paragraph(f'{pr_number}. {product.name}', style='List Bullet')
    doc.add_heading('Процессы', 2)
    for pr_index, product in enumerate(program.products.all()):
        product_number = pr_index + 1
        doc.add_heading(f'Продукт «{product_number}. {product.name}»', 3)
        for st_index, stage in enumerate(product.stages.all()):
            stage_number = st_index + 1
            doc.add_paragraph(f'Этап {product_number}.{stage_number}. {stage.name}', style='List Bullet')
            for proc_index, process in enumerate(stage.processes.all()):
                process_number = proc_index + 1
                doc.add_paragraph(f'Процесс {product_number}.{stage_number}.{process_number} {process.name}', style='List Bullet 2')

    doc.add_heading('Сводная таблица', level=1)

    # Получаем продукты и связанные данные
    products = program.products.prefetch_related('stages__processes').order_by('position')

    # Создаем таблицу
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Добавляем заголовки таблицы
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Продукты'
    header_cells[1].text = 'Этапы'
    header_cells[2].text = 'Процессы'

    # Заполняем таблицу данными
    for product in products:
        product_added = False  # Флаг для отслеживания добавления строки продукта

        for stage in product.stages.all():
            stage_added = False  # Флаг для отслеживания добавления строки этапа

            for process in stage.processes.all():
                # Добавляем строку для процесса
                row_cells = table.add_row().cells
                row_cells[2].text = process.name

                # Добавляем данные этапа только в первую строку для этапа
                if not stage_added:
                    row_cells[1].text = stage.name
                    stage_added = True

                # Добавляем данные продукта только в первую строку для продукта
                if not product_added:
                    row_cells[0].text = product.name
                    product_added = True

            # Если у этапа нет процессов, добавляем строку для этапа
            if not stage_added:
                row_cells = table.add_row().cells
                row_cells[1].text = stage.name
                if not product_added:
                    row_cells[0].text = product.name
                    product_added = True

        # Если у продукта нет этапов, добавляем строку для продукта
        if not product_added:
            row_cells = table.add_row().cells
            row_cells[0].text = product.name

    doc.add_heading('Компетентностный профиль', 1)

    doc.add_heading('Перечень знаний и умений', 2)
    for index, ability in enumerate(program.abilities.all()):
        ability_number = index + 1
        doc.add_paragraph(f'{ability_number}. {ability.name}', style='List Bullet')
        for kn_index, knowledge in enumerate(ability.knowledges.all()):
            knowledge_number = kn_index + 1
            doc.add_paragraph(f'{ability_number}.{knowledge_number} {knowledge.name}', style='List Bullet 2')

    doc.add_heading('Сопоставление процессов со знаниями и умениями', 2)

    for product in program.products.order_by('position'):
        for stage in product.stages.order_by('position'):
            for process in stage.processes.order_by('position'):
                doc.add_paragraph(f'Процесс: {process.name}', style='List Bullet')
                for ability in process.abilities.order_by('position'):
                    doc.add_paragraph(f'Умение: {ability.name}', style='List Bullet 2')
                    for knowledge in ability.knowledges.order_by('position'):
                        doc.add_paragraph(f'Знание: {knowledge.name}', style='List Bullet 3')

    # process_count_for_program = Program.objects.filter(id=program_id).aggregate(process_count=Count('products__stages__processes'))


    # table = doc.add_table(rows=process_count_for_program['process_count'], cols=3)
    # table.style='Table Grid'

    doc.add_heading('Оценочные материалы', level=1)
    doc.add_heading('Тестовые вопросы', level=2)
    knowledge_queryset = program.knowledges.annotate(question_count=Count('questions')).filter(question_count__gt=0)
    for knowledge in knowledge_queryset:
        # Добавляем уровень знания
        doc.add_paragraph(f'Знание: {knowledge.name}', style='List Bullet')

        for question in knowledge.questions.all():
            # Добавляем уровень вопросов
            doc.add_paragraph(f'Вопрос: {question.text} ({question.question_type.name})', style='List Bullet 2')

            is_sootv = question.question_type.name == 'Установление соответствия'
            for answer in question.answers.all():
                # Добавляем уровень ответов
                if is_sootv:
                    doc.add_paragraph(
                        f'Ответ: {answer.text} - {answer.text2}',
                        style='List Bullet 3'
                    )
                else:
                    doc.add_paragraph(
                        f'Ответ: {answer.text} {"(правильный)" if answer.is_correct else ""}',
                        style='List Bullet 3'
                    )

    abilities_with_tasks = Ability.objects.annotate(
        task_count=Count('tasks')
    ).filter(task_count__gt=0)

    doc.add_heading('Задания', level=2)
    for ability in abilities_with_tasks:
        doc.add_paragraph(f"Умение: {ability.name} (Количество заданий: {ability.task_count})", style='List Bullet')
        for task in ability.tasks.all():
            doc.add_paragraph(f"Задание: {task.name}", style='Normal')
            doc.add_paragraph(task.description, style='Normal')

    doc.add_heading('Дисциплины', 1)
    disciplines = program.disciplines.prefetch_related('processes').order_by('name')

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    header_cells[0].text = 'Дисциплины'
    header_cells[1].text = 'Процессы'

    # Заполняем таблицу данными
    for discipline in disciplines:
        row_cells = table.add_row().cells
        row_cells[0].text = discipline.name
        row_cells[1].text = ', '.join(process.name for process in discipline.processes.all())


    doc.add_heading('Дисциплины по семестрам', 1)
    for index, semester in enumerate(program.semesters.order_by('number'), start=1):
        doc.add_paragraph(f'Семестр №{semester.number}', style='List Bullet')

        for disc_index, discipline in enumerate(semester.disciplines.all(), start=1):
            doc.add_paragraph(f'{discipline.name}', style='List Bullet 2')



    # Сохраняем документ в BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)

    filename = quote('Дизайн программы.docx')
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename={filename}'
    response.write(output.read())

    return response


def add_multilevel_list(paragraph, text, level):
    run = paragraph.add_run(text)
    # Устанавливаем уровень списка
    paragraph_style = 'List Bullet'
    #paragraph.style = 'List Number ' + str(level)



def merge_cells(table, start_row, start_col, end_row, end_col):
    """Объединяет ячейки в таблице."""
    for row in range(start_row, end_row + 1):
        for col in range(start_col, end_col + 1):
            if row == start_row and col == start_col:
                continue
            cell = table.cell(row, col)
            cell._element.getparent().remove(cell._element)
    tc = table.cell(start_row, start_col)._element
    tc.set('rowSpan', str(end_row - start_row + 1))
    tc.set('colSpan', str(end_col - start_col + 1))