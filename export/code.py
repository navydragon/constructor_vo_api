import os
from docx import Document

# Укажите путь к директории проекта Django
project_path = r"C:\projects\constructor_vo_api\users"
# Укажите расширения файлов, которые хотите добавить в документ
extensions = {'.py',}

# Создаем новый документ Word
doc = Document()

# Проходим по всем файлам в проекте Django
for root, dirs, files in os.walk(project_path):
    for file in files:
        # Проверяем расширение файла
        if any(file.endswith(ext) for ext in extensions):
            file_path = os.path.join(root, file)
            doc.add_paragraph(f"Файл: {file_path}\n")

            # Читаем содержимое файла и добавляем его в документ
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                doc.add_paragraph(content)
                doc.add_paragraph("\n---\n")  # Разделитель между файлами

# Сохраняем документ
doc.save("Проект_Django_тексты.docx")
